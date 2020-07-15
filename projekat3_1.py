from tkinter import *
from tkinter.ttk import Combobox
import xlrd
from docx import Document
import os
from docx.shared import Inches
from tkcalendar import DateEntry

za_stampu = ""

root = Tk()   #pravljenje glavnog prozora programa

izbor_firmi = [   #niz u kome se cuvaju firme koje ce biti ponudjene kao izbor korisniku
    "firma 1",
    "firma 2",
    "firma 3"
]

dozvole = []  #niz za dozvole
dozvole_za_firmu = [] #niz za dozvole za firmu
varijabla_izbor_zemalja = StringVar() #varijabla za izbor zemalja

upisane_dozvole = []  #niz koji ce cuvati upisane dozvole

dozvole_izbor = Combobox(root,values=[1,2,3]) #pravljenje combobox-a za izbor dozvola
zemlje = Combobox(root,values=[1,2,3]) #pravljenje combobox-a za izbor zemalja
firme = Combobox() #pravljenje combobox-a za izbor firmi

def stampaj():          #funkcija za stampu
    global za_stampu    #globalna varijabla predvidjena za stampu

    if za_stampu == "": #ukoliko je varijabla za stampu prazna ispisuje se greska u labeli predvidjenoj za obavestavanje korisnika
        labela_greska_varijabla.set("niste snimili fajl za stampu")
    else:
        labela_greska_varijabla.set("stampanje")

def vrati_zemlje(): #funkcija kojom vracam imena svih zemalja
    global dozvole  #globalna varijabla sa imenima zemalja i njihovim dozvolama

    zemlje_lista = [] #pravim niz koji ce drzati imena zemalja kada ih procitam
    for i in dozvole:   #niz dozvole sadrzi recnike koji imaju zemlje i njihove dozvole
        #u petlji prolazim kroz niz iz recnika, citam ime zemlje i stavljam u novonapravljeni niz za zemlje
        zemlje_lista.append(i["zemlja"])
    return zemlje_lista #na kraju funkcija vraca niz koji sadrzi imena zemalja

def vrati_dozvole(zemlja):  #funkcija koja je predvidjena da vrati sve dozvole za jednu zemlju, zemlja se daje kao parametar
    global dozvole  #globalna varijabla sa imenima zemalja i njihovim dozvolama

    dozvole_lista = []    #pravim niz koji ce sadrzati sve dozvole za datu zemlju
    for i in dozvole:     #u petlji prolazim kroz niz koji sadrzi zemlje i dozvole
        if i["zemlja"] == zemlja:   #kada smo nasao u nizu zemlju koja mi je potrebna onda uzimam dozvole
            for j in i["dozvole"]:  #prolazim kroz dozvole za zemlju i upisujem ih u novonapravljeni niz
                dozvole_lista.append(j) #dodavanje dozvole
    return dozvole_lista    #vracam niz koji sadrzi dozvole za datu zemlju

def promeni_dozvole(index, value, op):  #funkcija predvidjena da promeni dozvole
    global dozvole_izbor     #globalna varijabla za izbor dozvola

    dozvole_izbor.destroy() #brisem padajuci meni sa dozvolama

    pom = vrati_dozvole(zemlje.get()) #uzimam nove dozvole za trenutnu zemlju

    dozvole_izbor = Combobox(frejm1, values = pom)    #pravim novi padajuci meni koji sadrzi dozvole za novoodabranu zemlju
    dozvole_izbor.set(pom[0])   #postavljam defoltnu vrednost
    dozvole_izbor.grid(row=3, column=0) #postavljam novi padajuci meni na mesto starog

def procitaj_dozvole(): #funkcija za citanje dozvola
    global dozvole  #globalna varijabla koja ce sadrzati spisak dozvola za svaku firmu
    #ova funkcija je predvidjena da napuni varijablu dozvolama koje ce mi kasnije biti potrebne

    #citanje dozvola iz xlsx fajla
    loc = ("zzz.xlsx")
    wb = xlrd.open_workbook(loc)
    sheet = wb.sheet_by_index(0)
    sheet.cell_value(0, 0)

    broj_zemalja = sheet.ncols  #brojim koliko kolona postoji, svaka zemlja ima svoju kolonu sa dozvolama

    for i in range(broj_zemalja):#u ovoj petlji prolazim kroz zemlje
        zemlja = sheet.col_values(i)[0]

        d = []  #niz koji ce sadrzati dozvole za trenutnu zemlju do koje se doslo u spoljnoj petlji
        for j in range(1, len(sheet.col_values(i))):#ova petlja prolazi kroz dozvole za trenutnu zemlju
            if sheet.col_values(i)[j] != '':    #provera da li dozvola postoji ili je prazna
                d.append(sheet.col_values(i)[j])    #dodajem trenutno procitanu dozvolu u niz koji sadrzi do sada procitane dozvole za zemlju

        dozvole.append({"zemlja": zemlja, "dozvole": d})    #posto sam procitao sve dozvole za trenutnu zemlju iz fajla onda dodajem u niz
        #novi recnik koji od atributa ima ime zemlje i sve njene ponudjene dozvole

def dodaj_opis_dozvole():
    global dozvole_za_firmu,upisane_dozvole

    if len(dozvole_za_firmu) == 0:#ukoliko su dozvole duzine 0 znaci da nema dodatih dozvola ili da su dozvole obrisane
        labela = StringVar()
        lb_nema = Label(root, textvariable = labela)
        lb_nema.grid(row=6, column=0, columnspan=2)
        labela.set("nema dodatih dozvola")
    else:
        labela = StringVar()
        lb = Label(root, textvariable = labela)
        lb.grid(row=len(dozvole_za_firmu)+6, column=0, columnspan=2)
        upisane_dozvole.append(lb)#dodajem labelu u listu upisanih dozvola, potrebno je voditi evideciju o ovoj listi da bi se mogla lako obrisati
        labela.set(str(len(dozvole_za_firmu))+". "+dozvole_za_firmu[len(dozvole_za_firmu)-1]["zemlja"]+"-"+dozvole_za_firmu[len(dozvole_za_firmu)-1]["dozvola"]+"    kom. "+dozvole_za_firmu[len(dozvole_za_firmu)-1]["broj"])

def ukloni_dozvole():   #uklanjam dozvole koje su upisane
    global upisane_dozvole,dozvole

    for i in upisane_dozvole:
        i.destroy()

    dozvole.clear()
    dozvole_za_firmu.clear()



def dodaj_dozvolu():#funkcija koja mi sluzi da doda dozvolu
    global dozvole_za_firmu
    odabrana_firma = firme.get()    #uzimam odabranu firmu iz padajuceg menija
    odabrana_zemlja = zemlje.get()  #uzimam odabranu zemlju iz padajuceg menija
    odabrana_dozvola = dozvole_izbor.get()  #uzimam odabranu dozvolu iz padajuceg menija
    odabran_broj = broj.get()   #uzimam odabran broj iz polja za unos broja

    if odabran_broj == "":    #provera da li je proj unet
        labela_greska_varijabla.set("unesite broj") #ispisuje mi gresku ako broj nije unet
    else:
        labela_greska_varijabla.set("") #ako je broj ispravno unesen potrebno je obrisati prethodno moguce ispise gresaka ako ih je bilo
        dozvole_za_firmu.append({"firma":odabrana_firma,"zemlja":odabrana_zemlja,"dozvola":odabrana_dozvola,"broj":odabran_broj})#dodajem dozvolu u niz dozvola za firme
        #dozvola je tipa recnik koja sadrzi ime zemlje ime dozvole i odabrani broj
        dodaj_opis_dozvole()

def upisi_u_word2(odabrana_firma,odabran_upis,odabran_datum):
    global  za_stampu,dozvole_za_firmu

    print("firma: "+odabrana_firma)
    print("mesto: "+odabran_upis)
    print("datum: "+odabran_datum)
    print("dozvole za zemlje:")
    br = 1
    for i in dozvole_za_firmu:
        print(str(br)+". "+i["zemlja"]+"-"+i["dozvola"]+"    kom. "+i["broj"])
        br+=1

def upisi_u_word(odabrana_firma,odabran_upis,odabran_datum):  #ova funkcija sluzi da bi se upisali podaci u word dokument
    global za_stampu,dozvole_za_firmu

    pom = odabran_datum.split("/")
    odabran_datum = pom[1] + "." + pom[0] + "." + pom[2]

    document = Document()

    document.add_heading('\n                                                            '
                           'Ministarstvo gradjevinarstva,\n'
                           '                                                            '
                           'saobracaja i infrastrukture')

    document.add_heading('\n\n\n                               Predmet: Zahtev za dodelu dozvola\n\n', level=1)

    br = 1
    for i in dozvole_za_firmu:
        document.add_paragraph(str(br) + ". " + i["zemlja"] + "-" + i["dozvola"] + "    kom. " + i["broj"])
        br += 1

    document.add_paragraph("\n\n\n\n"+odabran_upis+",\n"+odabran_datum+". godina")
    try:
        os.chdir("firme")
        os.chdir(odabrana_firma)
        os.chdir("dodela dozvola")
    except Exception as e:
        labela_greska_varijabla.set(str(e))

    try:
        document.save(odabran_datum+'.docx')
        labela_greska_varijabla.set("dokument snimljen")
    except Exception as e:
        labela_greska_varijabla.set(str(e))


def snimi_formu():  #ova funkcija snima formu
    odabrana_firma=firme.get()  #uzimam vrednost odabrane firme
    odabran_upis=upis.get() #uzimam vrednost za upis
    odabran_datum=upis_datum.get()  #uzimam vrednos odabranog datuma



    if len(dozvole_za_firmu) == 0:    #provera da li su unete dozvole
        labela_greska_varijabla.set("nema unetih dozvola")
    else:
        if odabran_upis == "":    #provera da li je unet upis
            labela_greska_varijabla.set("unesite nesto")
        else:
            labela_greska_varijabla.set("")
            if odabran_datum == "":   #provera da li je unet datum
                labela_greska_varijabla.set("unesite datum")
            else:
                upisi_u_word(odabrana_firma,odabran_upis,odabran_datum) #ako je sve u redu pozivam funkciju za upis u word

procitaj_dozvole()  #citam dozvole kako bi ih kasnije mogao koristiti


#ovde stavljam labelu koja govori korisniku da izabere firmu
labela_firma_varijabla=StringVar()
labela_firma = Label(root,textvariable=labela_firma_varijabla)
labela_firma.grid(row=0,column=0,columnspan=2)  #postavljamo popziciju
labela_firma_varijabla.set("izaberite firmu")

#kreiram padajuci meni koji mi daje izbor firmi moguce firme upisujem rucno, moze i da se promeni da se stavi lista koja vec drzi firme
firme=Combobox(root,values=["firma 1","firma 2","firma 3"])
firme.set("firma 1")    #postavljam defoltnu vrednost izbora
firme.grid(row=1,column=0,columnspan=2,padx=10,pady=10)  #postavljam popziciju


frejm1 = Frame(root, borderwidth = 1, relief = SUNKEN, padx=10,pady=10)
frejm1.grid(row=2,column=0)


#ovde stavljam labelu koja govori korisniku da izabere zemlju
labela_zemlja_varijabla=StringVar()
labela_zemlja=Label(frejm1,textvariable=labela_zemlja_varijabla)
labela_zemlja.grid(row=0,column=0)
labela_zemlja_varijabla.set("izaberite zemlju")

pom_zemlje=vrati_zemlje()   #ovde pravim varijablu u koju ce funkcija upisati sve zemlje koje se mogu izabrati

#kreiram padajuci meni koji ce mi omoguciti izbor zemalja, moguci izbor postavljam na vracene zemlje iz varijable pom_zemlje
zemlje=Combobox(frejm1,values=pom_zemlje,textvar=varijabla_izbor_zemalja)
zemlje.set(pom_zemlje[0])#postavljam defoltnu vrednost na prvu zemlju
zemlje.grid(row=1,column=0)
varijabla_izbor_zemalja.trace('w',promeni_dozvole)


#ovde stavljam labelu koja govori korisniku da izabere dozvole
labela_dozvola_varijabla=StringVar()
labela_dozvola=Label(frejm1,textvariable=labela_dozvola_varijabla)
labela_dozvola.grid(row=2,column=0)
labela_dozvola_varijabla.set("izaberite vrstu dozvole")

pom_dozvole=vrati_dozvole(pom_zemlje[0])#pravim varijablu u koju ce funkcija vratiti sve moguce dozvole za trenutno izabranu zemlju

dozvole_izbor=Combobox(frejm1,values=pom_dozvole)#pravim padajuci meni za izbor dozvola za trenutno odabranu zemlju
dozvole_izbor.set(pom_dozvole[0])#postavljam defoltnu vrednost na prvu dozvolu
dozvole_izbor.grid(row=3,column=0)

frejm2 = Frame(root, borderwidth = 1, relief = SUNKEN, padx=10,pady=10)
frejm2.grid(row=2,column=1)

#ovde stavljam labelu koja govori korisniku da izabere broj
labela_broj_varijabla=StringVar()
labela_broj=Label(frejm2,textvariable=labela_broj_varijabla)
labela_broj.grid(row=0,column=0)
labela_broj_varijabla.set("unesite broj")

#ovde pravim dugme za snimanje dozvola, kada se klikne poziva se funkcija dodaj dozvolu
snimi_novu_dozvolu = Button(frejm2, text="dodaj dozvolu",command=dodaj_dozvolu)
snimi_novu_dozvolu.grid(row=2,column=0,pady=10)

#polje za unos broja
broj=Entry(frejm2)
broj.grid(row=1,column=0)

frejm3 = Frame(root, borderwidth = 1, relief = SUNKEN, padx=10,pady=10)
frejm3.grid(row=3,column=0,columnspan=2)

labela_upis_varijabla=StringVar()
labela_upis=Label(frejm3,textvariable=labela_upis_varijabla)
labela_upis.grid(row=0,column=0)
labela_upis_varijabla.set("upisi nesto")

upis=Entry(frejm3)
upis.grid(row=1,column=0)

labela_upis_datum_varijabla=StringVar()
labela_upis_datum=Label(frejm3,textvariable=labela_upis_datum_varijabla)
labela_upis_datum.grid(row=0,column=1)
labela_upis_datum_varijabla.set("upisi datum")


upis_datum= DateEntry(frejm3, year=2020, month=7)
upis_datum.grid(row=1,column=1)

frejm4 = Frame(root, borderwidth = 1, relief = SUNKEN, padx=10,pady=10)
frejm4.grid(row=4,column=0,columnspan=2)

#pravim dugme za snimanje, kada se klikne poziva se funkcija snimi_formu
snimi = Button(frejm4, text="snimi kao",command=snimi_formu)
snimi.grid(row=0,column=0)

#pravim dugme za stampu, kada se klikne poziva se funkcija stampaj
stampaj = Button(frejm4, text="stampaj",command=stampaj)
stampaj.grid(row=0,column=1)

#pravimo labelu za gresku
labela_greska_varijabla=StringVar()
labela_greska=Label(root,textvariable=labela_greska_varijabla)
labela_greska.grid(row=5,column=0,columnspan=2)
labela_greska_varijabla.set("")

#pravim dugme za uklanjanje svih dozvola, kada se klikne poziva se funkcija ukloni_dozvole
ukloni = Button(root, text="ukloni dozvole",command=ukloni_dozvole)
ukloni.grid(row=6,column=0,columnspan=2)

if len(dozvole_za_firmu) == 0:
    labela = StringVar()
    lb = Label(root, textvariable=labela)
    lb.grid(row=7,column=0,columnspan=2)
    labela.set("nema dodatih dozvola")



root.mainloop()